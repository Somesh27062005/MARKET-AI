"""
agents/report_generation.py
Report Generation Agent — PDF + DOCX export
"""
import io, os, json
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors as rl_colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                 TableStyle, HRFlowable, PageBreak, KeepTogether)
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ── PDF Colors ───────────────────────────────────────────────────────────────
DARK = rl_colors.HexColor("#0f172a")     # Deep background slate
PRIMARY = rl_colors.HexColor("#3b82f6")  # Brand primary blue
SECONDARY = rl_colors.HexColor("#8b5cf6")# Secondary purple
ACCENT = rl_colors.HexColor("#60a5fa")   # Muted brand light blue
MUTED = rl_colors.HexColor("#94a3b8")    # Slate subtext gray
WHITE = rl_colors.white
CARD = rl_colors.HexColor("#1e293b")     # Cards background color
BORDER = rl_colors.HexColor("#334155")   # Grid boundaries border

# ── PDF Global Styles ─────────────────────────────────────────────────────────
styles = getSampleStyleSheet()

title_style = ParagraphStyle(
    "ReportTitle",
    parent=styles["Normal"],
    fontName="Helvetica-Bold",
    fontSize=26,
    leading=32,
    textColor=WHITE,
    alignment=TA_LEFT,
    spaceAfter=15
)

subtitle_style = ParagraphStyle(
    "ReportSubtitle",
    parent=styles["Normal"],
    fontName="Helvetica",
    fontSize=13,
    leading=18,
    textColor=ACCENT,
    spaceAfter=25
)

meta_style = ParagraphStyle(
    "ReportMeta",
    parent=styles["Normal"],
    fontName="Helvetica",
    fontSize=9,
    leading=14,
    textColor=MUTED,
    spaceAfter=30
)

h1_style = ParagraphStyle(
    "SectionHeader",
    parent=styles["Normal"],
    fontName="Helvetica-Bold",
    fontSize=15,
    leading=19,
    textColor=WHITE,
    spaceBefore=18,
    spaceAfter=10,
    keepWithNext=True
)

h2_style = ParagraphStyle(
    "SubSectionHeader",
    parent=styles["Normal"],
    fontName="Helvetica-Bold",
    fontSize=11.5,
    leading=15,
    textColor=PRIMARY,
    spaceBefore=12,
    spaceAfter=6,
    keepWithNext=True
)

body_style = ParagraphStyle(
    "BodyTextCustom",
    parent=styles["Normal"],
    fontName="Helvetica",
    fontSize=9,
    leading=14,
    textColor=rl_colors.HexColor("#cbd5e1"),
    spaceAfter=8
)

bullet_style = ParagraphStyle(
    "BulletTextCustom",
    parent=styles["Normal"],
    fontName="Helvetica",
    fontSize=9,
    leading=13.5,
    textColor=rl_colors.HexColor("#cbd5e1"),
    leftIndent=15,
    bulletIndent=5,
    spaceAfter=5
)

def _page_bg(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(DARK)
    canvas.rect(0, 0, A4[0], A4[1], fill=1, stroke=0)
    canvas.setFillColor(MUTED)
    canvas.setFont("Helvetica", 7)
    canvas.drawRightString(A4[0] - 30, 18, f"Page {doc.page}  ·  MarketAI Suite  ·  Confidential Platform Analytics")
    canvas.restoreState()

def _table(data, col_widths, header_color=PRIMARY):
    th = ParagraphStyle("TH_r", fontSize=8.5, fontName="Helvetica-Bold",
                         textColor=WHITE, alignment=TA_LEFT, leading=11)
    td = ParagraphStyle("TD_r", fontSize=8, fontName="Helvetica",
                         textColor=rl_colors.HexColor("#e2e8f0"), leading=11)
    rows = [[Paragraph(str(c).replace('\n', '<br/>'), th if i == 0 else td) for c in row]
            for i, row in enumerate(data)]
    t = Table(rows, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND",  (0, 0), (-1, 0),  header_color),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [CARD, DARK]),
        ("GRID",        (0, 0), (-1, -1), 0.5, BORDER),
        ("TOPPADDING",  (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("VALIGN",      (0, 0), (-1, -1), "TOP"),
    ]))
    return t

# ── PDF Exporter Main Entrypoint ──────────────────────────────────────────────
def generate_pdf(module: str, title: str, data: dict) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                             rightMargin=2*cm, leftMargin=2*cm,
                             topMargin=2*cm, bottomMargin=2*cm)
    W = A4[0] - 4*cm

    # ── Cover Page Layout ──
    story = [
        Spacer(1, 2.5*cm),
        Paragraph("MARKETAI SUITE PLATFORM BRIEF", ParagraphStyle("CoverBrand", fontName="Helvetica-Bold", fontSize=10, textColor=PRIMARY, spaceAfter=8)),
        Paragraph(title.upper(), title_style),
        HRFlowable(width=W, color=PRIMARY, thickness=3, spaceAfter=15),
        Paragraph(f"Module: {module.upper()} INTELLIGENCE SUMMARY", subtitle_style),
        Spacer(1, 3.5*cm),
        Paragraph(f"<b>Prepared by:</b> MarketAI Auto-Agent Orchestrator<br/>"
                  f"<b>Date Generated:</b> {datetime.now().strftime('%B %d, %Y')}<br/>"
                  f"<b>Pipeline Architecture:</b> LangGraph + LangChain RAG<br/>"
                  f"<b>Target Scope:</b> {data.get('company_name', 'Global Market Operations')}<br/>"
                  f"<b>Classification:</b> Proprietary & Confidential", meta_style),
        PageBreak(),
    ]

    # ── Module-specific content generation ──
    if module == "market":
        _pdf_market(story, data, W)
    elif module == "campaign":
        _pdf_campaign(story, data, W)
    elif module == "pitch":
        _pdf_pitch(story, data, W)
    elif module == "lead":
        _pdf_lead(story, data, W)
    elif module == "insights":
        _pdf_insights(story, data, W)

    doc.build(story, onFirstPage=_page_bg, onLaterPages=_page_bg)
    return buf.getvalue()

def _pdf_market(story, d, W):
    story.append(Paragraph("Executive Summary", h1_style))
    story.append(Paragraph(d.get("executive_summary", "No executive summary available."), body_style))
    story.append(Spacer(1, 10))

    ms = d.get("market_size", {})
    if ms:
        story.append(Paragraph("Market Projections & Growth CAGR", h1_style))
        data = [
            ["Metric Detail", "Value Projection"],
            ["Current Estimated Size", ms.get("current", "N/A")],
            ["Projected Market Size", ms.get("projected", "N/A")],
            ["Compound Annual Growth Rate (CAGR)", ms.get("cagr", "N/A")]
        ]
        story.append(_table(data, [W * 0.6, W * 0.4]))
        story.append(Spacer(1, 10))

    competitors = d.get("competitors", [])
    if competitors:
        story.append(Paragraph("Competitive Threat Matrix", h1_style))
        rows = [["Competitor Name", "Strengths", "Weaknesses", "Market Position", "Threat Level"]]
        for c in competitors[:6]:
            rows.append([
                c.get("name",""),
                c.get("strengths","")[:90],
                c.get("weaknesses","")[:90],
                c.get("market_position",""),
                c.get("threat_level","")
            ])
        story.append(_table(rows, [W * 0.18, W * 0.32, W * 0.32, W * 0.10, W * 0.08]))
        story.append(Spacer(1, 10))

    swot = d.get("swot", {})
    if swot:
        story.append(Paragraph("SWOT Analysis Matrix", h1_style))
        rows = [
            ["Strengths", "Weaknesses"],
            [
                "\n".join(f"• {x}" for x in swot.get("strengths", [])),
                "\n".join(f"• {x}" for x in swot.get("weaknesses", []))
            ],
            ["Opportunities", "Threats"],
            [
                "\n".join(f"• {x}" for x in swot.get("opportunities", [])),
                "\n".join(f"• {x}" for x in swot.get("threats", []))
            ]
        ]
        story.append(_table(rows, [W / 2, W / 2]))
        story.append(Spacer(1, 10))

    pestel = d.get("pestel", {})
    if pestel:
        story.append(Paragraph("PESTEL Macro Environment Factors", h1_style))
        rows = [["Factor Dimension", "Assessment & Impact Analysis"]]
        for factor in ["political", "economic", "social", "technological", "environmental", "legal"]:
            rows.append([factor.title(), pestel.get(factor, "N/A")])
        story.append(_table(rows, [W * 0.25, W * 0.75]))
        story.append(Spacer(1, 10))

    advertising = d.get("advertising_analysis", [])
    if advertising:
        story.append(Paragraph("Advertising Channel Benchmarks", h1_style))
        rows = [["Channel/Platform", "CPM/CPC", "Creative Strategy", "Message Angle", "Efficiency"]]
        for ad in advertising:
            rows.append([
                ad.get("channel", ""),
                ad.get("cpm_cpc_benchmark", ""),
                ad.get("creative_strategy", ""),
                ad.get("message_angle", ""),
                ad.get("ad_spend_efficiency", "")
            ])
        story.append(_table(rows, [W * 0.20, W * 0.18, W * 0.25, W * 0.25, W * 0.12]))
        story.append(Spacer(1, 10))

    postures = d.get("positioning_postures", [])
    if postures:
        story.append(Paragraph("Competitive Positioning Postures", h1_style))
        rows = [["Brand/Competitor", "Market Role", "Pricing Posture", "Innovation Posture", "Acquisition Strategy"]]
        for p in postures:
            rows.append([
                p.get("brand_name", ""),
                p.get("market_role", ""),
                p.get("pricing_posture", ""),
                p.get("innovation_posture", ""),
                p.get("customer_acquisition_posture", "")
            ])
        story.append(_table(rows, [W * 0.22, W * 0.18, W * 0.18, W * 0.20, W * 0.22]))

def _pdf_campaign(story, d, W):
    story.append(Paragraph(f"Campaign Strategy: {d.get('campaign_name', 'Unnamed Campaign')}", h1_style))
    story.append(Paragraph("Strategic Objectives", h2_style))
    for obj in d.get("objectives", []):
        story.append(Paragraph(f"• {obj}", bullet_style))
    story.append(Spacer(1, 10))

    persona = d.get("persona", {})
    if persona:
        story.append(Paragraph("Target Buyer Persona", h2_style))
        p_name = persona.get("name", "Target Segment")
        p_desc = f"<b>Role:</b> {persona.get('role', 'N/A')} | <b>Age Range:</b> {persona.get('age_range', 'N/A')}<br/>" \
                 f"<b>Industry Focus:</b> {persona.get('industry', 'N/A')}<br/>" \
                 f"<b>Core Pain Points:</b> {', '.join(persona.get('pain_points', []))}<br/>" \
                 f"<b>Primary Goals:</b> {', '.join(persona.get('goals', []))}"
        story.append(Paragraph(f"<b>Profile: {p_name}</b><br/>{p_desc}", body_style))
        story.append(Spacer(1, 10))

    ad_copies = d.get("ad_copies", [])
    if ad_copies:
        story.append(Paragraph("Creative Ad Copy Variations", h2_style))
        for copy in ad_copies[:3]:
            story.append(Paragraph(f"<b>Platform: {copy.get('platform','')}</b>", h2_style))
            story.append(Paragraph(f"<b>Headline:</b> {copy.get('headline','')}", body_style))
            story.append(Paragraph(f"<b>Body Copy:</b> {copy.get('body','')}", body_style))
            story.append(Paragraph(f"<b>Call to Action (CTA):</b> {copy.get('cta','')}", body_style))
            story.append(Spacer(1, 5))

    budget = d.get("budget_allocation", [])
    if budget:
        story.append(Paragraph("Budget Allocation & Rationale", h2_style))
        rows = [["Channel/Platform", "Allocation %", "Strategic Rationale"]]
        for b in budget:
            rows.append([b.get("channel",""), f"{b.get('percent',0)}%", b.get("rationale","")])
        story.append(_table(rows, [W * 0.25, W * 0.15, W * 0.60]))
        story.append(Spacer(1, 10))

    kpis = d.get("kpis", [])
    if kpis:
        story.append(Paragraph("KPI & Success Measurement Metrics", h2_style))
        rows = [["KPI Metric", "Target Goal", "Measurement Methodology"]]
        for k in kpis:
            rows.append([k.get("metric",""), k.get("target",""), k.get("measurement","")])
        story.append(_table(rows, [W * 0.35, W * 0.25, W * 0.40]))
        story.append(Spacer(1, 10))

    calendar = d.get("calendar", [])
    if calendar:
        story.append(Paragraph("Campaign Timeline Roadmap", h2_style))
        rows = [["Week", "Weekly Focus/Theme", "Tasks Checklist"]]
        for item in calendar[:6]:
            rows.append([
                f"Week {item.get('week', '')}",
                item.get("theme", ""),
                "\n".join(f"- {t}" for t in item.get("tasks", []))
        story.append(_table(rows, [W * 0.15, W * 0.35, W * 0.50]))

    social_posts = d.get("social_media_posts", [])
    if social_posts:
        story.append(Spacer(1, 10))
        story.append(Paragraph("Social Media Posts", h2_style))
        for post in social_posts[:6]:
            platform = post.get("platform", "Social Media")
            copy = post.get("copy", "")
            media = post.get("media_suggestion", "")
            story.append(Paragraph(f"<b>Platform:</b> {platform}", body_style))
            story.append(Paragraph(f"<b>Content:</b> {copy}", body_style))
            if media:
                story.append(Paragraph(f"<b>Media Suggestion:</b> {media}", bullet_style))
            story.append(Spacer(1, 5))

def _pdf_pitch(story, d, W):
    story.append(Paragraph("Sales Pitch & Prepared Objections", h1_style))
    story.append(Paragraph("Elevator Pitch", h2_style))
    story.append(Paragraph(d.get("elevator_pitch", "N/A"), body_style))
    story.append(Spacer(1, 10))

    vp = d.get("value_proposition", {})
    if vp:
        story.append(Paragraph(f"Value Proposition: {vp.get('headline','')}", h2_style))
        for pt in vp.get("points", []):
            story.append(Paragraph(f"• {pt}", bullet_style))
        story.append(Spacer(1, 10))

    roi = d.get("roi_argument", {})
    if roi:
        story.append(Paragraph(f"ROI Financial Arguments", h2_style))
        story.append(Paragraph(f"<b>Business Return Case:</b> {roi.get('headline','')}<br/><b>Calculation Details:</b> {roi.get('calculation','')}<br/><b>Timeframe:</b> {roi.get('timeframe','')}", body_style))
        story.append(Spacer(1, 10))

    objections = d.get("objection_handling", [])
    if objections:
        story.append(Paragraph("Objection Handling Workbook", h2_style))
        for item in objections[:4]:
            story.append(Paragraph(f"<b>Objection:</b> {item.get('objection','')}", body_style))
            story.append(Paragraph(f"<b>Response Strategy:</b> {item.get('response','')}", body_style))
            story.append(Spacer(1, 5))

    email = d.get("email_template", {})
    if email:
        story.append(Paragraph("Outbound Sales Email Template", h2_style))
        story.append(Paragraph(f"<b>Subject:</b> {email.get('subject','')}", body_style))
        story.append(Paragraph(email.get("body","").replace("\n", "<br/>"), body_style))
        story.append(Spacer(1, 10))

    linkedin = d.get("linkedin_template", {})
    if linkedin:
        story.append(Paragraph("LinkedIn Outreach Script", h2_style))
        story.append(Paragraph(f"<b>Connection Note:</b> {linkedin.get('connection_note','')}<br/><b>Follow Up:</b> {linkedin.get('follow_up','')}", body_style))

def _pdf_lead(story, d, W):
    score = d.get("lead_score", 0)
    temp = d.get("temperature", "Warm")
    prob = d.get("conversion_probability", 0)
    story.append(Paragraph("Predictive Lead Scoring Assessment", h1_style))
    story.append(Paragraph(f"<b>Lead Score:</b> {score}/100  ·  <b>Category:</b> {temp}  ·  <b>Conversion Probability:</b> {prob}%", h2_style))
    story.append(Paragraph(d.get("qualification_summary", ""), body_style))
    story.append(Spacer(1, 10))

    bant = d.get("bant", {})
    if bant:
        story.append(Paragraph("BANT Matrix Breakdown", h2_style))
        rows = [["Dimension", "Score", "Qualifying Assessment", "Evidence Details"]]
        for key in ["budget", "authority", "need", "timeline"]:
            item = bant.get(key, {})
            rows.append([
                key.title(),
                f"{item.get('score', 0)}/25",
                item.get("assessment", "N/A"),
                item.get("evidence", "N/A")[:100]
            ])
        story.append(_table(rows, [W * 0.18, W * 0.12, W * 0.35, W * 0.35]))
        story.append(Spacer(1, 10))

    intent = d.get("buying_intent", {})
    if intent:
        story.append(Paragraph(f"Buying Intent Signals (Aggregate: {intent.get('score', 0)}/100)", h2_style))
        for sig in intent.get("signals", []):
            story.append(Paragraph(f"• <b>[{sig.get('strength', 'Low')}]</b> {sig.get('signal', 'N/A')}", bullet_style))
        story.append(Spacer(1, 10))

    risks = d.get("risk_factors", [])
    if risks:
        story.append(Paragraph("Identified Deal Risk Factors", h2_style))
        rows = [["Risk Factor Details", "Severity", "Mitigation Strategy"]]
        for r in risks:
            rows.append([r.get("risk", ""), r.get("impact", ""), r.get("mitigation", "")])
        story.append(_table(rows, [W * 0.35, W * 0.15, W * 0.50]))
        story.append(Spacer(1, 10))

    story.append(Paragraph("Next Steps & Actions", h2_style))
    story.append(Paragraph(f"<b>Next Best Action:</b> {d.get('next_best_action', 'N/A')}", body_style))

def _pdf_insights(story, d, W):
    story.append(Paragraph("Business Optimization Advisory", h1_style))
    story.append(Paragraph(f"<b>Opportunity Score:</b> {d.get('opportunity_score', 0)}/100", h2_style))
    story.append(Paragraph(d.get("executive_summary", ""), body_style))
    story.append(Spacer(1, 10))

    challenges = d.get("current_challenges", [])
    if challenges:
        story.append(Paragraph("Identified Challenges & Friction Points", h2_style))
        rows = [["Blocker Challenge", "Severity", "Impact Analysis"]]
        for c in challenges:
            rows.append([c.get("challenge", ""), c.get("severity", ""), c.get("impact", "")])
        story.append(_table(rows, [W * 0.35, W * 0.15, W * 0.50]))
        story.append(Spacer(1, 10))

    recs = d.get("strategic_recommendations", [])
    if recs:
        story.append(Paragraph("Strategic Action Recommendations", h2_style))
        rows = [["Action Recommendation", "Priority", "Impact Level", "Effort Rating"]]
        for r in recs:
            rows.append([r.get("recommendation", ""), r.get("priority", ""), r.get("impact", ""), r.get("effort", "")])
        story.append(_table(rows, [W * 0.45, W * 0.15, W * 0.20, W * 0.20]))
        story.append(Spacer(1, 10))

    revenue = d.get("revenue_opportunities", [])
    cost = d.get("cost_optimization", [])
    if revenue or cost:
        story.append(Paragraph("Financial Action Items", h2_style))
        if revenue:
            story.append(Paragraph("<b>Revenue Growth Areas:</b>", body_style))
            for item in revenue:
                story.append(Paragraph(f"• {item.get('source', '')} (Potential: {item.get('potential', '')} | Timeline: {item.get('timeline', '')})", bullet_style))
        if cost:
            story.append(Paragraph("<b>Cost Saving Areas:</b>", body_style))
            for item in cost:
                story.append(Paragraph(f"• Optimize {item.get('area', '')} (Potential: {item.get('potential_savings', '')} | Step: {item.get('action', '')})", bullet_style))
        story.append(Spacer(1, 10))

    for phase, label in [("plan_30_day","30-Day Plan"), ("plan_60_day","60-Day Plan"), ("plan_90_day","90-Day Plan")]:
        story.append(PageBreak())
        story.append(Paragraph(f"Execution Roadmap: {label}", h1_style))
        story.append(Spacer(1, 10))
        rows = [["Roadmap Action", "Owner", "Key Success Metric"]]
        for item in d.get(phase, []):
            rows.append([item.get("action",""), item.get("owner",""), item.get("success_metric","")])
        story.append(_table(rows, [W * 0.45, W * 0.20, W * 0.35]))

# ── DOCX Helper XML Shading & Border Methods ───────────────────────────────
def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_callout_borders(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    # Left border: thick brand highlight line
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '36')  # 4.5 pt width
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), color_hex)
    tcBorders.append(left)
    
    # Clear all other borders
    for border_name in ['top', 'bottom', 'right']:
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'nil')
        tcBorders.append(b)
    tcPr.append(tcBorders)

# ── DOCX Formatting Helper Methods ──────────────────────────────────────────
def add_styled_paragraph(doc, text, style_name="Normal", font_name="Arial", size_pt=10, color_rgb=None, bold=False, italic=False, align=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align:
        p.alignment = align
    
    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb
    return p

def add_bullet_point(doc, text, bold_prefix="", font_name="Arial", size_pt=9.5, space_after=4):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(space_after)
    if bold_prefix:
        r1 = p.add_run(str(bold_prefix))
        r1.bold = True
        r1.font.name = font_name
        r1.font.size = Pt(size_pt)
    r2 = p.add_run(str(text))
    r2.font.name = font_name
    r2.font.size = Pt(size_pt)
    return p

def add_callout_box(doc, text, title="EXECUTIVE ADVISORY HIGHLIGHT", border_color_hex="3B82F6", bg_color_hex="F8FAFC"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, bg_color_hex)
    set_callout_borders(cell, border_color_hex)
    
    # Padding twips inside table cell (144 = 7.2pt, 216 = 10.8pt)
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m in ['top', 'bottom', 'left', 'right']:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), '144' if m in ['top', 'bottom'] else '216')
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    
    r_title = p.add_run(f"{title}\n")
    r_title.bold = True
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(9)
    r_title.font.color.rgb = RGBColor(59, 130, 246)
    
    r_body = p.add_run(str(text))
    r_body.italic = True
    r_body.font.name = 'Arial'
    r_body.font.size = Pt(9.5)
    r_body.font.color.rgb = RGBColor(71, 85, 105)
    
    # Spacing line after table callout
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(4)

def format_docx_table(table, col_widths, header_bg="1F2937", alt_bg="F9FAFB"):
    for i, row in enumerate(table.rows):
        # Prevent row split across pages
        trPr = row._tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)
        
        # Format individual cells
        for j, cell in enumerate(row.cells):
            cell.width = col_widths[j]
            if i == 0:
                set_cell_shading(cell, header_bg)
                # White header text styling
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.space_before = Pt(4)
                    for r in p.runs:
                        r.bold = True
                        r.font.color.rgb = RGBColor(255, 255, 255)
                        r.font.size = Pt(9.5)
            else:
                if i % 2 == 1:
                    set_cell_shading(cell, alt_bg)
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.space_before = Pt(4)
                    for r in p.runs:
                        r.font.size = Pt(9)
                        r.font.color.rgb = RGBColor(55, 65, 81)

# ── Word (DOCX) Exporter Main Entrypoint ──────────────────────────────────────
def generate_docx(module: str, title: str, data: dict) -> bytes:
    doc = Document()
    
    # Page setup margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Cover Page Block ──
    add_styled_paragraph(doc, "MARKETAI SUITE INTELLIGENCE BRIEFING", size_pt=10, color_rgb=RGBColor(59, 130, 246), bold=True, space_before=72, space_after=10)
    add_styled_paragraph(doc, title.upper(), size_pt=28, color_rgb=RGBColor(15, 23, 42), bold=True, space_after=6)
    
    # Border line
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_after = Pt(12)
    run_line = p_line.add_run("─" * 45)
    run_line.bold = True
    run_line.font.color.rgb = RGBColor(59, 130, 246)
    
    add_styled_paragraph(doc, f"Module Classification: {module.upper()} STRATEGY DOCUMENT", size_pt=12, color_rgb=RGBColor(100, 116, 139), italic=True, space_after=120)
    
    # Metadata Block
    add_styled_paragraph(doc, "DOCUMENT INFORMATION", size_pt=9.5, color_rgb=RGBColor(15, 23, 42), bold=True, space_after=4)
    add_styled_paragraph(doc, f"Generated Date: {datetime.now().strftime('%B %d, %Y')}\n"
                              f"Company Context: {data.get('company_name', 'Global Operations')}\n"
                              f"Orchestrated By: LangGraph Agent Node Synthesis\n"
                              f"Document Security: Confidential & Restricted", size_pt=9, color_rgb=RGBColor(100, 116, 139), space_after=12)
    
    doc.add_page_break()

    # ── Module Content Generation ──
    if module == "market":
        _docx_market(doc, data)
    elif module == "campaign":
        _docx_campaign(doc, data)
    elif module == "pitch":
        _docx_pitch(doc, data)
    elif module == "lead":
        _docx_lead(doc, data)
    elif module == "insights":
        _docx_insights(doc, data)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def _docx_market(doc, d):
    add_styled_paragraph(doc, "Executive Summary", size_pt=16, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=12, space_after=8)
    add_callout_box(doc, d.get("executive_summary", "N/A"), title="MARKET STRATEGY SUMMARY")
    
    ms = d.get("market_size", {})
    if ms:
        add_styled_paragraph(doc, "Market Sizing & Projections", size_pt=14, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=12, space_after=6)
        table = doc.add_table(rows=4, cols=2)
        table.rows[0].cells[0].text = "Sizing Vector"
        table.rows[0].cells[1].text = "Value Estimation"
        
        metrics = [
            ("Current Market Valuation", ms.get("current", "N/A")),
            ("Projected Market Valuation", ms.get("projected", "N/A")),
            ("Compound Annual Growth Rate (CAGR)", ms.get("cagr", "N/A"))
        ]
        for idx, (metric, val) in enumerate(metrics, start=1):
            table.rows[idx].cells[0].text = metric
            table.rows[idx].cells[1].text = val
        format_docx_table(table, [Inches(3.5), Inches(3.0)])
        doc.add_paragraph()

    competitors = d.get("competitors", [])
    if competitors:
        add_styled_paragraph(doc, "Competitive Landscape Analysis", size_pt=14, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=12, space_after=6)
        table = doc.add_table(rows=1, cols=5)
        table.rows[0].cells[0].text = "Competitor"
        table.rows[0].cells[1].text = "Strengths"
        table.rows[0].cells[2].text = "Weaknesses"
        table.rows[0].cells[3].text = "Market Position"
        table.rows[0].cells[4].text = "Threat"
        
        for c in competitors[:6]:
            row = table.add_row()
            row.cells[0].text = c.get("name","")
            row.cells[1].text = c.get("strengths","")[:90]
            row.cells[2].text = c.get("weaknesses","")[:90]
            row.cells[3].text = c.get("market_position","")
            row.cells[4].text = c.get("threat_level","")
            
        format_docx_table(table, [Inches(1.2), Inches(1.8), Inches(1.8), Inches(1.0), Inches(0.7)])
        doc.add_paragraph()

    swot = d.get("swot", {})
    if swot:
        add_styled_paragraph(doc, "SWOT Analysis Matrix", size_pt=14, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=12, space_after=6)
        table = doc.add_table(rows=4, cols=2)
        table.rows[0].cells[0].text = "Strengths (Internal)"
        table.rows[0].cells[1].text = "Weaknesses (Internal)"
        table.rows[1].cells[0].text = "\n".join(f"• {x}" for x in swot.get("strengths", []))
        table.rows[1].cells[1].text = "\n".join(f"• {x}" for x in swot.get("weaknesses", []))
        table.rows[2].cells[0].text = "Opportunities (External)"
        table.rows[2].cells[1].text = "Threats (External)"
        table.rows[3].cells[0].text = "\n".join(f"• {x}" for x in swot.get("opportunities", []))
        table.rows[3].cells[1].text = "\n".join(f"• {x}" for x in swot.get("threats", []))
        format_docx_table(table, [Inches(3.25), Inches(3.25)])
        doc.add_paragraph()

    pestel = d.get("pestel", {})
    if pestel:
        add_styled_paragraph(doc, "PESTEL Macro Analysis", size_pt=14, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=12, space_after=6)
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Macro Factor"
        table.rows[0].cells[1].text = "Strategic Impact Assessment"
        for key in ["political", "economic", "social", "technological", "environmental", "legal"]:
            row = table.add_row()
            row.cells[0].text = key.title()
            row.cells[1].text = pestel.get(key, "N/A")
        format_docx_table(table, [Inches(1.8), Inches(4.7)])
        doc.add_paragraph()

    advertising = d.get("advertising_analysis", [])
    if advertising:
        add_styled_paragraph(doc, "Advertising Channel Benchmarks", size_pt=14, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=12, space_after=6)
        table = doc.add_table(rows=1, cols=5)
        table.rows[0].cells[0].text = "Channel/Platform"
        table.rows[0].cells[1].text = "CPM/CPC Benchmark"
        table.rows[0].cells[2].text = "Creative Strategy"
        table.rows[0].cells[3].text = "Message Angle"
        table.rows[0].cells[4].text = "Efficiency"
        for ad in advertising:
            row = table.add_row()
            row.cells[0].text = ad.get("channel", "")
            row.cells[1].text = ad.get("cpm_cpc_benchmark", "")
            row.cells[2].text = ad.get("creative_strategy", "")
            row.cells[3].text = ad.get("message_angle", "")
            row.cells[4].text = ad.get("ad_spend_efficiency", "")
        format_docx_table(table, [Inches(1.3), Inches(1.2), Inches(1.5), Inches(1.5), Inches(1.0)])
        doc.add_paragraph()

    postures = d.get("positioning_postures", [])
    if postures:
        add_styled_paragraph(doc, "Competitive Positioning Postures", size_pt=14, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=12, space_after=6)
        table = doc.add_table(rows=1, cols=5)
        table.rows[0].cells[0].text = "Competitor Brand"
        table.rows[0].cells[1].text = "Market Role"
        table.rows[0].cells[2].text = "Pricing Posture"
        table.rows[0].cells[3].text = "Innovation Posture"
        table.rows[0].cells[4].text = "Acquisition Vector"
        for p in postures:
            row = table.add_row()
            row.cells[0].text = p.get("brand_name", "")
            row.cells[1].text = p.get("market_role", "")
            row.cells[2].text = p.get("pricing_posture", "")
            row.cells[3].text = p.get("innovation_posture", "")
            row.cells[4].text = p.get("customer_acquisition_posture", "")
        format_docx_table(table, [Inches(1.4), Inches(1.1), Inches(1.2), Inches(1.3), Inches(1.5)])
        doc.add_paragraph()

def _docx_campaign(doc, d):
    add_styled_paragraph(doc, f"Campaign Strategy Overview: {d.get('campaign_name', 'Strategy')}", size_pt=16, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=12, space_after=8)
    
    add_styled_paragraph(doc, "Campaign Objectives", size_pt=12, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=8, space_after=4)
    for obj in d.get("objectives", []):
        add_bullet_point(doc, obj)
        
    persona = d.get("persona", {})
    if persona:
        add_styled_paragraph(doc, "Target Segment Buyer Persona", size_pt=12, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=10, space_after=4)
        p_desc = f"Representative Persona: {persona.get('name', 'N/A')}\n" \
                 f"Role Profile: {persona.get('role', 'N/A')}  |  Demographics: {persona.get('age_range', 'N/A')}  |  Industry Scope: {persona.get('industry', 'N/A')}\n" \
                 f"Pain Points: {', '.join(persona.get('pain_points', []))}\n" \
                 f"Segment Goals: {', '.join(persona.get('goals', []))}"
        add_callout_box(doc, p_desc, title="BUYER PERSONA ASSIGNMENT")

    ad_copies = d.get("ad_copies", [])
    if ad_copies:
        add_styled_paragraph(doc, "Platform Ad Copies Variations", size_pt=12, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=10, space_after=6)
        for copy in ad_copies[:3]:
            add_styled_paragraph(doc, f"Channel Type: {copy.get('platform', '')}", size_pt=10, color_rgb=RGBColor(59, 130, 246), bold=True)
            add_styled_paragraph(doc, f"Headline: {copy.get('headline', '')}", size_pt=9.5, italic=True)
            add_styled_paragraph(doc, f"Body Copy: {copy.get('body', '')}", size_pt=9.5)
            add_styled_paragraph(doc, f"Call-to-Action: {copy.get('cta', '')}", size_pt=9.5, bold=True, space_after=8)

    budget = d.get("budget_allocation", [])
    if budget:
        add_styled_paragraph(doc, "Marketing Budget Split Allocation", size_pt=12, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=10, space_after=6)
        table = doc.add_table(rows=1, cols=3)
        table.rows[0].cells[0].text = "Marketing Channel"
        table.rows[0].cells[1].text = "Allocation %"
        table.rows[0].cells[2].text = "Strategic Rationale"
        for b in budget:
            row = table.add_row()
            row.cells[0].text = b.get("channel", "")
            row.cells[1].text = f"{b.get('percent', 0)}%"
            row.cells[2].text = b.get("rationale", "")
        format_docx_table(table, [Inches(1.8), Inches(1.2), Inches(3.5)])
        doc.add_paragraph()

    calendar = d.get("calendar", [])
    if calendar:
        add_styled_paragraph(doc, "Campaign Timeline Calendar Roadmap", size_pt=12, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=10, space_after=6)
        table = doc.add_table(rows=1, cols=3)
        table.rows[0].cells[0].text = "Timeline"
        table.rows[0].cells[1].text = "Weekly Focus Theme"
        table.rows[0].cells[2].text = "Tasks Checklist"
        for item in calendar[:6]:
            row = table.add_row()
            row.cells[0].text = f"Week {item.get('week', '')}"
            row.cells[1].text = item.get("theme", "")
            row.cells[2].text = ", ".join(item.get("tasks", []))
        format_docx_table(table, [Inches(1.0), Inches(2.2), Inches(3.3)])
        doc.add_paragraph()

    social_posts = d.get("social_media_posts", [])
    if social_posts:
        add_styled_paragraph(doc, "Social Media Posts Campaign", size_pt=12, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=10, space_after=6)
        for post in social_posts[:6]:
            platform = post.get("platform", "Social Media")
            copy = post.get("copy", "")
            media = post.get("media_suggestion", "")
            add_styled_paragraph(doc, f"Platform: {platform}", size_pt=10, color_rgb=RGBColor(59, 130, 246), bold=True, space_before=4)
            add_styled_paragraph(doc, copy, size_pt=9.5)
            if media:
                add_styled_paragraph(doc, f"Media Recommendation: {media}", size_pt=9.5, italic=True, space_after=8)

def _docx_pitch(doc, d):
    add_styled_paragraph(doc, "Sales Pitch Handbook & Objection Scripts", size_pt=16, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=12, space_after=8)
    
    add_styled_paragraph(doc, "Elevator Pitch", size_pt=12, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=8, space_after=4)
    add_callout_box(doc, d.get("elevator_pitch", "N/A"), title="30-SECOND ELEVATOR PITCH")

    vp = d.get("value_proposition", {})
    if vp:
        add_styled_paragraph(doc, f"Core Value Proposition: {vp.get('headline','')}", size_pt=12, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=10, space_after=4)
        for pt in vp.get("points", []):
            add_bullet_point(doc, pt)
        doc.add_paragraph()

    roi = d.get("roi_argument", {})
    if roi:
        add_styled_paragraph(doc, "ROI & Commercial Payback Argument", size_pt=12, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=10, space_after=4)
        add_styled_paragraph(doc, f"Argument: {roi.get('headline','')}", size_pt=10, bold=True)
        add_styled_paragraph(doc, f"Financial ROI Calculation: {roi.get('calculation','')}", size_pt=9.5)
        add_styled_paragraph(doc, f"Realization Timeframe: {roi.get('timeframe','')}", size_pt=9.5, space_after=10)

    objections = d.get("objection_handling", [])
    if objections:
        add_styled_paragraph(doc, "Objections Playbook Guidelines", size_pt=12, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=10, space_after=6)
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Expected Objection"
        table.rows[0].cells[1].text = "Recommended Sales Response Script"
        for item in objections[:4]:
            row = table.add_row()
            row.cells[0].text = item.get("objection", "")
            row.cells[1].text = item.get("response", "")
        format_docx_table(table, [Inches(2.5), Inches(4.0)])
        doc.add_paragraph()

    email = d.get("email_template", {})
    if email:
        add_styled_paragraph(doc, "Outbound Sales Email Template", size_pt=12, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=10, space_after=4)
        email_body = f"Subject: {email.get('subject', '')}\n\n{email.get('body', '')}"
        add_callout_box(doc, email_body, title="OUTBOUND COLD EMAIL TEMPLATE", border_color_hex="8B5CF6")

def _docx_lead(doc, d):
    score = d.get("lead_score", 0)
    temp = d.get("temperature", "Warm")
    prob = d.get("conversion_probability", 0)
    
    add_styled_paragraph(doc, "Predictive Lead Scoring Assessment", size_pt=16, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=12, space_after=8)
    
    summary = f"Qualitative Qualification Summary:\n{d.get('qualification_summary','')}"
    add_callout_box(doc, summary, title=f"LEAD SCORE: {score}/100 ({temp} | CONVERSION PROBABILITY: {prob}%)", border_color_hex="10B981")

    bant = d.get("bant", {})
    if bant:
        add_styled_paragraph(doc, "BANT Qualification Matrix Analysis", size_pt=12, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=10, space_after=6)
        table = doc.add_table(rows=1, cols=4)
        table.rows[0].cells[0].text = "BANT Dimension"
        table.rows[0].cells[1].text = "Score"
        table.rows[0].cells[2].text = "Assessment Details"
        table.rows[0].cells[3].text = "Evidence Verified"
        for key in ["budget", "authority", "need", "timeline"]:
            item = bant.get(key, {})
            row = table.add_row()
            row.cells[0].text = key.title()
            row.cells[1].text = f"{item.get('score', 0)}/25"
            row.cells[2].text = item.get("assessment", "N/A")
            row.cells[3].text = item.get("evidence", "N/A")[:90]
        format_docx_table(table, [Inches(1.5), Inches(0.8), Inches(2.2), Inches(2.0)])
        doc.add_paragraph()

    risks = d.get("risk_factors", [])
    if risks:
        add_styled_paragraph(doc, "Deal Risk Assessments", size_pt=12, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=10, space_after=6)
        table = doc.add_table(rows=1, cols=3)
        table.rows[0].cells[0].text = "Identified Threat Risk"
        table.rows[0].cells[1].text = "Severity Rating"
        table.rows[0].cells[2].text = "Mitigation Path Action"
        for r in risks:
            row = table.add_row()
            row.cells[0].text = r.get("risk", "")
            row.cells[1].text = r.get("impact", "")
            row.cells[2].text = r.get("mitigation", "")
        format_docx_table(table, [Inches(2.5), Inches(1.0), Inches(3.0)])
        doc.add_paragraph()

    add_styled_paragraph(doc, f"Recommended Next Action: {d.get('next_best_action', 'N/A')}", size_pt=11, bold=True, color_rgb=RGBColor(16, 185, 129))

def _docx_insights(doc, d):
    add_styled_paragraph(doc, "Business Insights Strategy Consultation", size_pt=16, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=12, space_after=8)
    
    add_callout_box(doc, d.get("executive_summary", ""), title=f"EXECUTIVE CONSULTATION (OPPORTUNITY SCORE: {d.get('opportunity_score', 0)}/100)", border_color_hex="F59E0B")

    challenges = d.get("current_challenges", [])
    if challenges:
        add_styled_paragraph(doc, "Identified Challenges & Operational Friction", size_pt=12, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=10, space_after=6)
        table = doc.add_table(rows=1, cols=3)
        table.rows[0].cells[0].text = "Friction Challenge"
        table.rows[0].cells[1].text = "Severity"
        table.rows[0].cells[2].text = "Impact Description"
        for c in challenges:
            row = table.add_row()
            row.cells[0].text = c.get("challenge", "")
            row.cells[1].text = c.get("severity", "")
            row.cells[2].text = c.get("impact", "")
        format_docx_table(table, [Inches(2.2), Inches(1.0), Inches(3.3)])
        doc.add_paragraph()

    recs = d.get("strategic_recommendations", [])
    if recs:
        add_styled_paragraph(doc, "Strategic Advisory Recommendations", size_pt=12, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=10, space_after=6)
        table = doc.add_table(rows=1, cols=4)
        table.rows[0].cells[0].text = "Strategic Path"
        table.rows[0].cells[1].text = "Priority"
        table.rows[0].cells[2].text = "Impact"
        table.rows[0].cells[3].text = "Implementation Effort"
        for r in recs:
            row = table.add_row()
            row.cells[0].text = r.get("recommendation", "")
            row.cells[1].text = r.get("priority", "")
            row.cells[2].text = r.get("impact", "")
            row.cells[3].text = r.get("effort", "")
        format_docx_table(table, [Inches(3.2), Inches(1.0), Inches(1.1), Inches(1.2)])
        doc.add_paragraph()

    # 30-60-90 Day Execution Roadmap Tables
    for phase, label in [("plan_30_day","30-Day Transition Roadmap"), ("plan_60_day","60-Day Mid-Term Roadmap"), ("plan_90_day","90-Day Execution Roadmap")]:
        add_styled_paragraph(doc, label, size_pt=13, color_rgb=RGBColor(31, 41, 55), bold=True, space_before=12, space_after=6)
        table = doc.add_table(rows=1, cols=3)
        table.rows[0].cells[0].text = "Execution Action"
        table.rows[0].cells[1].text = "Accountable Owner"
        table.rows[0].cells[2].text = "Key Success Metric"
        for item in d.get(phase, []):
            row = table.add_row()
            row.cells[0].text = item.get("action","")
            row.cells[1].text = item.get("owner","")
            row.cells[2].text = item.get("success_metric","")
        format_docx_table(table, [Inches(3.2), Inches(1.3), Inches(2.0)])
        doc.add_paragraph()
